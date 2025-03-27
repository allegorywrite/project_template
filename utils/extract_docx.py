#!/usr/bin/env python3
"""
DOCXファイルからテキストを抽出するスクリプト
"""

import sys
import os
from pathlib import Path
import docx

def extract_text_from_docx(docx_path, output_dir=None):
    """
    DOCXファイルからテキストを抽出する
    
    Args:
        docx_path (str or Path): DOCXファイルのパス
        output_dir (str or Path, optional): 出力ディレクトリ
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        print(f"エラー: ファイルが見つかりません: {docx_path}")
        return
    
    if output_dir is None:
        output_dir = docx_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    # 出力ファイル名
    markdown_file = output_dir / f"{docx_path.stem}_text.md"
    
    try:
        # DOCXファイルを開く
        doc = docx.Document(docx_path)
        
        # テキストを抽出
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # テーブルからテキストを抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        # 結果を保存
        print(f"結果を保存しています: {markdown_file}")
        with open(markdown_file, "w", encoding="utf-8") as f:
            f.write("\n\n".join(full_text))
        print(f"結果を保存しました: {markdown_file}")
        
        return "\n\n".join(full_text)
        
    except Exception as e:
        print(f"エラー: {e}")
        import traceback
        traceback.print_exc()
        return None

def main():
    if len(sys.argv) < 2:
        print("使用方法: python extract_docx.py <docx_path> [output_dir]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    extract_text_from_docx(docx_path, output_dir)

if __name__ == "__main__":
    main()
